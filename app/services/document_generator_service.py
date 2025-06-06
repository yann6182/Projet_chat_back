import os
import logging
import uuid
from datetime import datetime
from typing import Optional, Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentGeneratorService:
    def __init__(self, output_dir: str = "data/generated_docs"):
        """
        Initialise le service de génération de documents.
        
        Args:
            output_dir: Le répertoire où seront stockés les documents générés
        """
        self.output_dir = output_dir
        
        # Créer le répertoire de sortie s'il n'existe pas
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            logger.info(f"Répertoire de sortie créé: {output_dir}")
    
    def generate_pdf(self, 
                     title: str, 
                     content: str, 
                     metadata: Optional[Dict[str, Any]] = None,
                     sources: Optional[List[str]] = None) -> str:
        """
        Génère un document PDF à partir du contenu fourni.
        
        Args:
            title: Le titre du document
            content: Le contenu principal du document
            metadata: Métadonnées supplémentaires (optionnel)
            sources: Liste des sources utilisées dans la réponse (optionnel)
            
        Returns:
            Le chemin vers le fichier PDF généré
        """
        try:
            # Générer un nom de fichier unique
            filename = f"{uuid.uuid4().hex}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(self.output_dir, filename)
            
            # Créer le document PDF
            doc = SimpleDocTemplate(
                file_path,
                pagesize=letter,
                rightMargin=72, leftMargin=72,
                topMargin=72, bottomMargin=72
            )            # Styles
            styles = getSampleStyleSheet()
            
            # Modifier les styles existants au lieu de créer de nouveaux
            title_style = styles['Title']
            title_style.alignment = 1  # Centré
            title_style.fontSize = 16
            title_style.spaceAfter = 24
            
            # Modifier le style Normal existant
            normal_style = styles['Normal']
            normal_style.fontSize = 11
            normal_style.spaceAfter = 6
            
            # Ajouter un style pour les sous-titres
            styles.add(ParagraphStyle(
                name='Subtitle', 
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12
            ))            # Créer un style personnalisé pour les sources
            styles.add(ParagraphStyle(
                name='DocSources',  # Renommer pour éviter tout conflit potentiel 
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=3,
                leftIndent=20
            ))
            
            # Contenu du document
            elements = []
            
            # Titre
            elements.append(Paragraph(title, styles['Title']))
            
            # Date de génération
            date_str = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            elements.append(Paragraph(date_str, styles['Normal']))
            elements.append(Spacer(1, 24))
            
            # Métadonnées (si fournies)
            if metadata:
                elements.append(Paragraph("Informations", styles['Subtitle']))
                for key, value in metadata.items():
                    elements.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
                elements.append(Spacer(1, 12))
            
            # Contenu principal
            elements.append(Paragraph("Contenu", styles['Subtitle']))
            
            # Traiter le contenu ligne par ligne pour préserver le formatage
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    elements.append(Paragraph(paragraph, styles['Normal']))
            
            # Ajouter les sources si disponibles
            if sources and len(sources) > 0:
                elements.append(Spacer(1, 12))
                elements.append(Paragraph("Sources", styles['Subtitle']))
                for source in sources:
                    elements.append(Paragraph(f"• {source}", styles['DocSources']))
            
            # Générer le document
            doc.build(elements)
            logger.info(f"Document PDF généré avec succès: {file_path}")
            
            return file_path
        
        except Exception as e:
            logger.error(f"Erreur lors de la génération du document PDF: {str(e)}")
            raise
    
    def generate_word(self, 
                      title: str, 
                      content: str, 
                      metadata: Optional[Dict[str, Any]] = None,
                      sources: Optional[List[str]] = None) -> str:
        """
        Génère un document Word (DOCX) à partir du contenu fourni.
        
        Args:
            title: Le titre du document
            content: Le contenu principal du document
            metadata: Métadonnées supplémentaires (optionnel)
            sources: Liste des sources utilisées dans la réponse (optionnel)
            
        Returns:
            Le chemin vers le fichier DOCX généré
        """
        try:
            # Générer un nom de fichier unique
            filename = f"{uuid.uuid4().hex}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = os.path.join(self.output_dir, filename)
            
            # Créer le document Word
            doc = Document()
            
            # Titre
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date de génération
            date_str = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            date_paragraph = doc.add_paragraph()
            date_paragraph.add_run(date_str)
            
            doc.add_paragraph()  # Espacement
            
            # Métadonnées (si fournies)
            if metadata:
                doc.add_heading("Informations", level=2)
                for key, value in metadata.items():
                    meta_paragraph = doc.add_paragraph()
                    meta_paragraph.add_run(f"{key}: ").bold = True
                    meta_paragraph.add_run(str(value))
                
                doc.add_paragraph()  # Espacement
            
            # Contenu principal
            doc.add_heading("Contenu", level=2)
            
            # Traiter le contenu ligne par ligne pour préserver le formatage
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Ajouter les sources si disponibles
            if sources and len(sources) > 0:
                doc.add_paragraph()  # Espacement
                doc.add_heading("Sources", level=2)
                
                for source in sources:
                    source_paragraph = doc.add_paragraph()
                    source_paragraph.add_run(f"• {source}")
                    source_paragraph.paragraph_format.left_indent = Inches(0.25)
            
            # Sauvegarder le document
            doc.save(file_path)
            logger.info(f"Document Word généré avec succès: {file_path}")
            
            return file_path
        
        except Exception as e:
            logger.error(f"Erreur lors de la génération du document Word: {str(e)}")
            raise
